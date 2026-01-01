#!/usr/bin/env python3
"""
Frost & Noel KDP DOCX Converter

Converts markdown episode files to KDP-ready DOCX with:
- Title page with series name, episode title, subtitle, author (centered, page break)
- Copyright page (centered, 10pt font, page break)
- Chapter content (Heading 1 with page break, left justified with 0.5" first line indent)
- Back matter: Prequel CTA (centered), About Author (left justified), Mailing List (centered)

Usage:
    python convert_to_kdp_docx.py <episode_folder> [output_path]

Example:
    python convert_to_kdp_docx.py ep02_two_turtle_doves
    python convert_to_kdp_docx.py ep03_three_french_hens "C:/Output/Book3.docx"
"""

import os
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# =============================================================================
# CONFIGURATION
# =============================================================================

SERIES_NAME = "Frost & Noel"
SERIES_SUBTITLE = "Tinsel Town"
AUTHOR_NAME = "Marc Fist"
PUBLISHER_NAME = "Two-Fisted Reads"
COPYRIGHT_YEAR = "2025"

# Links
PREQUEL_LINK = "https://dl.bookfunnel.com/cbv3qk3sof"
MAILING_LIST_LINK = "https://marc-fist.kit.com/18679abba6"
AMAZON_AUTHOR_LINK = "https://www.amazon.com/stores/Marc-Fist/author/B0DY4PR11Y"

# Font sizes
BODY_FONT_SIZE = Pt(11)
COPYRIGHT_FONT_SIZE = Pt(10)
HEADING_FONT_SIZE = Pt(14)

# Indentation
FIRST_LINE_INDENT = Inches(0.5)

# Episode metadata (episode_folder_name -> (title, subtitle))
EPISODE_METADATA = {
    "ep01_partridge_in_a_panic": ("Partridge in a Panic", "Grumpy Meets Sunshine"),
    "ep02_two_turtle_doves": ("Two Turtle Doves", "First Kiss"),
    "ep03_three_french_hens": ("Three French Hens", "First Time"),
    "ep04_four_calling_birds": ("Four Calling Birds", "Going Public"),
    "ep05_five_gold_rings": ("Five Gold Rings", "The Promise"),
    "ep06_six_geese_a_laying": ("Six Geese a-Laying", "Found Family"),
    "ep07_seven_swans_a_swimming": ("Seven Swans a-Swimming", "Taking Care"),
    "ep08_eight_maids_a_milking": ("Eight Maids a-Milking", "First I Love You"),
    "ep09_nine_ladies_dancing": ("Nine Ladies Dancing", "The Rupture"),
    "ep10_ten_lords_a_leaping": ("Ten Lords a-Leaping", "Rock Bottom"),
    "ep11_eleven_pipers_piping": ("Eleven Pipers Piping", "Repair Begins"),
    "ep12_twelve_drummers_drumming": ("Twelve Drummers Drumming", "Happily Ever After"),
}

# =============================================================================
# FRONT MATTER TEMPLATES
# =============================================================================

COPYRIGHT_TEXT = f"""Copyright © {COPYRIGHT_YEAR} {AUTHOR_NAME}

All rights reserved

The characters and events portrayed in this book are fictitious. Any similarity to real persons, living or dead, is coincidental and not intended by the author.

No part of this book may be reproduced, or stored in a retrieval system, or transmitted in any form or by any means, electronic, mechanical, photocopying, recording, or otherwise, without express written permission of the publisher."""

# =============================================================================
# BACK MATTER TEMPLATES
# =============================================================================

PREQUEL_CTA_HEADING = "Before the Twelve Days begin..."
PREQUEL_CTA_SUBTITLE = "Meet Eben Frost the night before."
PREQUEL_CTA_BODY = """One phone call. One contract. One very stubborn rule about not getting attached.

Eben Frost drives through the night to a town that shouldn't exist, for a job he shouldn't have taken, to help a man he hasn't met yet.

This is the moment before everything changes."""
PREQUEL_CTA_BUTTON = "Send me the prequel"

PREQUEL_INFO = """Frost and Noel
Tinsel Town
The Night Before
One Grump, Too Much Christmas"""

ABOUT_AUTHOR_HEADING = "About The Author"
ABOUT_AUTHOR_TEXT = f"""{AUTHOR_NAME}

I'm Marc Fist, and I'm not afraid to get my hands dirty. I write about desire, connection, and all the messy, beautiful things that happen when two people finally stop pretending they don't want each other.

Visit my author page on Amazon to find more of my books."""

MAILING_LIST_HEADING = "Mailing List"
MAILING_LIST_TEXT = """Want to be the first to know?

Then subscribe to The Fist!

Join the mailing list"""

# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def add_page_break(doc):
    """Add a page break to the document."""
    doc.add_page_break()


def add_paragraph_centered(doc, text, font_size=BODY_FONT_SIZE, bold=False):
    """Add a centered paragraph."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.size = font_size
    run.font.bold = bold
    return para


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Blue underline style
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return paragraph


def add_paragraph_with_link(doc, text, url, font_size=BODY_FONT_SIZE, centered=True):
    """Add a paragraph with a hyperlink."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    add_hyperlink(para, text, url)
    return para


def add_paragraph_left_justified(doc, text, font_size=BODY_FONT_SIZE, first_line_indent=None):
    """Add a left-justified paragraph with optional first line indent."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    if first_line_indent:
        para.paragraph_format.first_line_indent = first_line_indent
    
    run = para.add_run(text)
    run.font.size = font_size
    return para


def add_heading_with_page_break(doc, text, level=1):
    """Add a heading with a page break before it."""
    add_page_break(doc)
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading


# =============================================================================
# CONVERSION FUNCTIONS
# =============================================================================

def find_manuscript(episode_folder: Path) -> Path:
    """Find the best manuscript file in the episode folder."""
    # Priority: COMPLETE_*.md > rev folders > draft folders
    
    # Check for compiled manuscript
    for f in episode_folder.glob("COMPLETE_*.md"):
        return f
    
    # Check for rev folders (highest number first)
    rev_folders = sorted(episode_folder.glob("rev_*"), reverse=True)
    for rev_folder in rev_folders:
        chapters = sorted(rev_folder.glob("chapter_*.md"))
        if chapters:
            return rev_folder
    
    # Check draft folders
    draft_folders = sorted(episode_folder.glob("draft_*"), reverse=True)
    for draft_folder in draft_folders:
        chapters = sorted(draft_folder.glob("chapter_*.md"))
        if chapters:
            return draft_folder
    
    raise FileNotFoundError(f"No manuscript found in {episode_folder}")


def load_chapters_from_folder(folder: Path) -> list[tuple[str, str, str, str]]:
    """Load chapters from a folder of chapter_XX.md files.
    Returns list of (chapter_num, subtitle, pov, body) tuples.
    """
    chapters = []
    for chapter_file in sorted(folder.glob("chapter_*.md")):
        content = chapter_file.read_text(encoding="utf-8")
        chapter_num = str(len(chapters) + 1)
        subtitle = ""
        pov = ""
        
        # Extract chapter title if present (matches both # and ##)
        title_match = re.search(r"^#+\s*Chapter\s*(\d+)[:\s]*(.*)$", content, re.MULTILINE | re.IGNORECASE)
        if title_match:
            chapter_num = title_match.group(1)
            subtitle = title_match.group(2).strip()
        
        # Remove markdown headers and clean content
        body = re.sub(r"^#+\s*.*$", "", content, flags=re.MULTILINE)
        body = body.strip()
        
        # Check for POV indicator at start of body (e.g., **NOEL** or NOEL)
        pov_match = re.match(r"^\*\*([A-Z]+)\*\*\s*", body)
        if pov_match:
            pov = pov_match.group(1)
            body = body[pov_match.end():].strip()
        else:
            # Also check for all-caps line at start
            pov_match = re.match(r"^([A-Z]{2,})\s*\n", body)
            if pov_match:
                pov = pov_match.group(1)
                body = body[pov_match.end():].strip()
        
        chapters.append((chapter_num, subtitle, pov, body))
    
    return chapters


def load_chapters_from_compiled(file: Path) -> list[tuple[str, str, str, str]]:
    """Load chapters from a single compiled manuscript.
    Returns list of (chapter_num, subtitle, pov, body) tuples.
    """
    content = file.read_text(encoding="utf-8")
    
    # Split on chapter headings (matches # or ##)
    chapter_pattern = r"^#+\s*Chapter\s*(\d+)[:\s]*(.*)$"
    parts = re.split(chapter_pattern, content, flags=re.MULTILINE | re.IGNORECASE)
    
    chapters = []
    # parts[0] is everything before first chapter
    for i in range(1, len(parts), 3):
        if i + 2 < len(parts):
            chapter_num = parts[i]
            subtitle = parts[i + 1].strip()
            body = parts[i + 2].strip()
            pov = ""
            
            # Check for POV indicator at start of body (e.g., **NOEL** or NOEL)
            pov_match = re.match(r"^\*\*([A-Z]+)\*\*\s*", body)
            if pov_match:
                pov = pov_match.group(1)
                body = body[pov_match.end():].strip()
            else:
                # Also check for all-caps line at start
                pov_match = re.match(r"^([A-Z]{2,})\s*\n", body)
                if pov_match:
                    pov = pov_match.group(1)
                    body = body[pov_match.end():].strip()
            
            chapters.append((chapter_num, subtitle, pov, body))
        elif i + 1 < len(parts):
            chapter_num = parts[i]
            body = parts[i + 1].strip()
            chapters.append((chapter_num, "", "", body))
    
    return chapters


def clean_markdown(text: str) -> str:
    """Clean markdown formatting for DOCX."""
    # Remove bold/italic markers
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    
    # Fix smart quotes and dashes
    text = text.replace("—", "—")  # em-dash
    text = text.replace("'", "'")
    text = text.replace("'", "'")
    text = text.replace(""", '"')
    text = text.replace(""", '"')
    
    # Remove horizontal rules
    text = re.sub(r"^---+$", "", text, flags=re.MULTILINE)
    
    return text.strip()


def create_kdp_docx(episode_folder: str, output_path: str = None) -> str:
    """Create KDP-ready DOCX from episode folder."""
    
    episode_path = Path(episode_folder)
    if not episode_path.exists():
        # Try relative to Frost & Noel project
        project_root = Path(r"c:\Users\Ann\Documents\novel stuff\ANTIGRAVITY\Agent Manager\projects\Frost & Noel")
        episode_path = project_root / episode_folder
    
    if not episode_path.exists():
        raise FileNotFoundError(f"Episode folder not found: {episode_folder}")
    
    # Get episode metadata
    folder_name = episode_path.name
    if folder_name in EPISODE_METADATA:
        episode_title, episode_subtitle = EPISODE_METADATA[folder_name]
    else:
        episode_title = folder_name.replace("_", " ").title()
        episode_subtitle = ""
    
    # Find and load manuscript
    manuscript = find_manuscript(episode_path)
    if manuscript.is_dir():
        chapters = load_chapters_from_folder(manuscript)
    else:
        chapters = load_chapters_from_compiled(manuscript)
    
    if not chapters:
        raise ValueError(f"No chapters found in {manuscript}")
    
    print(f"Found {len(chapters)} chapters")
    
    # Create document
    doc = Document()
    
    # ==========================================================================
    # TITLE PAGE (Centered)
    # ==========================================================================
    
    # Add some vertical spacing at top
    for _ in range(5):
        doc.add_paragraph()
    
    add_paragraph_centered(doc, f"{SERIES_NAME} - {SERIES_SUBTITLE}", font_size=Pt(18), bold=True)
    add_paragraph_centered(doc, episode_title, font_size=Pt(24), bold=True)
    add_paragraph_centered(doc, episode_subtitle, font_size=Pt(14))
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_paragraph_centered(doc, AUTHOR_NAME, font_size=Pt(12))
    add_paragraph_centered(doc, PUBLISHER_NAME, font_size=Pt(10))
    
    # ==========================================================================
    # COPYRIGHT PAGE (Centered, 10pt font)
    # ==========================================================================
    
    add_page_break(doc)
    
    # Copyright heading
    heading = doc.add_heading("Copyright", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Vertical spacing
    
    for line in COPYRIGHT_TEXT.split("\n\n"):
        add_paragraph_centered(doc, line, font_size=COPYRIGHT_FONT_SIZE)
    
    # ==========================================================================
    # CHAPTER CONTENT (Page break before each, left justified with indent)
    # ==========================================================================
    
    for i, (chapter_num, subtitle, pov, chapter_body) in enumerate(chapters, 1):
        # Page break and chapter heading
        add_page_break(doc)
        
        # Chapter number as heading (e.g., "Chapter 1")
        heading = doc.add_heading(f"Chapter {chapter_num}", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle on separate line (centered, italic)
        if subtitle:
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.size = Pt(12)
            subtitle_run.font.italic = True
        
        # POV indicator on separate line (centered, italic)
        if pov:
            pov_para = doc.add_paragraph()
            pov_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pov_run = pov_para.add_run(pov)
            pov_run.font.size = Pt(11)
            pov_run.font.italic = True
        
        doc.add_paragraph()  # Space after heading block
        
        # Add chapter body paragraphs (left justified with first line indent)
        body = clean_markdown(chapter_body)
        for para_text in body.split("\n\n"):
            para_text = para_text.strip()
            if para_text:
                add_paragraph_left_justified(
                    doc, 
                    para_text, 
                    font_size=BODY_FONT_SIZE,
                    first_line_indent=FIRST_LINE_INDENT
                )
    
    # ==========================================================================
    # BACK MATTER: PREQUEL CTA (Centered)
    # ==========================================================================
    
    add_page_break(doc)
    
    for _ in range(3):
        doc.add_paragraph()
    
    add_paragraph_centered(doc, PREQUEL_CTA_HEADING, font_size=Pt(16), bold=True)
    add_paragraph_centered(doc, PREQUEL_CTA_SUBTITLE, font_size=Pt(12))
    
    doc.add_paragraph()
    
    for line in PREQUEL_CTA_BODY.split("\n\n"):
        add_paragraph_centered(doc, line, font_size=BODY_FONT_SIZE)
    
    doc.add_paragraph()
    add_paragraph_with_link(doc, PREQUEL_CTA_BUTTON, PREQUEL_LINK)
    
    doc.add_paragraph()
    for line in PREQUEL_INFO.split("\n"):
        add_paragraph_centered(doc, line, font_size=Pt(10))
    
    # ==========================================================================
    # BACK MATTER: ABOUT AUTHOR (Left justified)
    # ==========================================================================
    
    add_page_break(doc)
    
    heading = doc.add_heading(ABOUT_AUTHOR_HEADING, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    for line in ABOUT_AUTHOR_TEXT.split("\n\n"):
        if "Visit my author page" in line:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.add_run("Visit my ")
            add_hyperlink(para, "author page on Amazon", AMAZON_AUTHOR_LINK)
            para.add_run(" to find more of my books.")
        else:
            add_paragraph_left_justified(doc, line, font_size=BODY_FONT_SIZE)
    
    # ==========================================================================
    # BACK MATTER: MAILING LIST (Centered)
    # ==========================================================================
    
    add_page_break(doc)
    
    for _ in range(3):
        doc.add_paragraph()
    
    heading = doc.add_heading(MAILING_LIST_HEADING, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    for line in MAILING_LIST_TEXT.split("\n\n"):
        if "Join the mailing list" in line:
            add_paragraph_with_link(doc, "Join the mailing list", MAILING_LIST_LINK)
        else:
            add_paragraph_centered(doc, line, font_size=BODY_FONT_SIZE)
    
    # ==========================================================================
    # SAVE
    # ==========================================================================
    
    if output_path:
        out = Path(output_path)
    else:
        out = episode_path / f"kdp-{folder_name}.docx"
    
    doc.save(str(out))
    print(f"Created: {out}")
    return str(out)


# =============================================================================
# CLI
# =============================================================================

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    
    episode = sys.argv[1]
    output = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        result = create_kdp_docx(episode, output)
        print(f"\nSuccess! KDP-ready DOCX created at:\n{result}")
    except Exception as e:
        print(f"ERROR: {e}")
        sys.exit(1)
